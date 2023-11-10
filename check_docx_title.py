import os  
from docx import Document
from tqdm import tqdm 
# 参考
# https://python-docx.readthedocs.io/en/latest/
# pip install python-docx
# https://zhuanlan.zhihu.com/p/377134370
# http://www.docpe.com/word/combine-word.aspx


output_folder = "data/output"
input = "data/05 产业上中下游价格指数周度报告/产业上中下游价格指数报告_49篇"
# input = "data/06 高频产销存数据分析快报/高频产销存数据分析快报_49篇"
if not os.path.exists(output_folder):  
    os.makedirs(output_folder)  
    print(f"Path {output_folder} created.")  
else:  
    print(f"Path {output_folder} already exists.")  

document=Document()

# 获取文件列表并排序  
file_list = os.listdir(input)  
file_list.sort()
# print(file_list)
# 创建一个新文件，如果文件已经存在则会被覆盖  

with open("output.txt", "w") as file:  
    file.write(f"输出结果：\n")  
    for i, filename in enumerate(file_list):
        if filename.endswith(".docx"):  
            doc_path = os.path.join(input, filename)  
            doc = Document(doc_path)  
            first_paragraph = doc.paragraphs[0].text  
            second_paragraph = doc.paragraphs[1].text  
            first_paragraph = f'{first_paragraph}（第{i+1}期）'  
            print(filename, first_paragraph, second_paragraph)  
            doc.paragraphs[0].text = first_paragraph 
            output_filename = os.path.join(output_folder, filename)  
            doc.save(output_filename)  
            with open('output_title.txt', 'a+') as f:  
                f.write(f"{filename}: {first_paragraph}")  
                f.write(f"{second_paragraph}\n")
            # print(f"文档已修改并保存到：{output_filename}")