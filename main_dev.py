import os
import re
from docx import Document
from docxcompose.composer import Composer
from tqdm import tqdm

def extract_numbers(s):
    return [int(match) for match in re.findall(r'\d+', s)]

def remove_second_paragraph(doc):
    if len(doc.paragraphs) >= 2:
        second_paragraph = doc.paragraphs[1]
        second_paragraph.clear()

        for run in second_paragraph.runs:
            new_run = doc.paragraphs[0].add_run(run.text, run.style)
            new_run.bold = run.bold
            # 其他样式属性也可以适当设置

        doc._element.body.remove(second_paragraph._element)

def process_word_documents(input_folder, output_folder, use_filename_as_title=False):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
        print(f"Path {output_folder} created.")
    else:
        print(f"Path {output_folder} already exists.")

    # Check if input_folder exists
    if not os.path.exists(input_folder):
        print(f"Input folder '{input_folder}' does not exist. Exiting.")
        return

    # 获取文件列表并排序
    file_list = os.listdir(input_folder)
    sorted_file_list = sorted(file_list, key=lambda x: extract_numbers(x))  # 使用自定义排序函数
    print("Sorted File List:")
    print("\n".join(sorted_file_list))

    with open("output.txt", "w") as file:
        file.write(f"输出结果：\n")
        for i, filename in enumerate(sorted_file_list):
            if filename.endswith(".docx"):
                doc_path = os.path.join(input_folder, filename)
                doc = Document(doc_path)
                first_paragraph = doc.paragraphs[0].text
                second_paragraph = doc.paragraphs[1].text

                if use_filename_as_title:
                    # 使用文件名作为标题
                    first_paragraph = os.path.splitext(filename)[0]
                else:
                    # 第二段加到第一段中
                    first_paragraph = f'{doc.paragraphs[0].text} {doc.paragraphs[1].text}'
                    remove_second_paragraph(doc)
                
                print(filename, first_paragraph, second_paragraph, end="\n\n")
                doc.paragraphs[0].text = first_paragraph
                output_filename = os.path.join(output_folder, filename)
                doc.save(output_filename)
                with open('output_title.txt', 'a+') as f:
                    f.write(f"{filename}: {first_paragraph}")
                    f.write(f"{second_paragraph}\n")

def merge_and_process_documents(input_folder, output_folder, perform_modification=True, use_filename_as_title=False):
    if perform_modification:
        process_word_documents(input_folder, output_folder, use_filename_as_title)

    # 获取修改后的文档列表
    modified_files = os.listdir(output_folder)

    # 构建修改后文档的完整路径列表
    modified_files_paths = [os.path.join(output_folder, file) for file in modified_files]

    # 执行文档合并
    merge_doc(source_file_path_list=modified_files_paths, target_folder='/Users/harvin/code/docx_merge/data/merged', input_folder=input_folder)

def merge_doc(source_file_path_list, target_folder, input_folder):
    page_break_doc = Document()
    page_break_doc.add_page_break()

    # 对文件路径进行排序
    sorted_source_files = sorted(source_file_path_list, key=lambda x: extract_numbers(x))
    print("Sorted Source Files:")
    print("\n".join(sorted_source_files))

    # 获取输入文件夹的最后一级文件夹的名字
    input_folder_name = os.path.basename(os.path.normpath(input_folder))

    # 构建目标文件的完整路径
    target_file_prefix = os.path.join(target_folder, f"{input_folder_name}合并")
    target_file_extension = ".docx"
    target_file_path = None

    target_doc = Document(sorted_source_files[0])
    target_composer = Composer(target_doc)
    for i in range(len(sorted_source_files)):
        # 跳过第一个作为模板的文件
        if i == 0:
            continue
        # 填充分页符文档
        target_composer.append(page_break_doc)
        # 拼接文档内容
        f = sorted_source_files[i]
        target_composer.append(Document(f))

    # 保存目标文档
    target_file_path = f"{target_file_prefix}（1-{len(sorted_source_files)}）{target_file_extension}"
    target_composer.save(target_file_path)

if __name__ == '__main__':
    print(os.getcwd())

    # Use full paths or adjust as needed
    input_folder = '/Users/harvin/code/docx_merge/data/5.5 （25-24）工程机械指数分析旬度报告'
    output_folder = '/Users/harvin/code/docx_merge/data/output/5.5 （25-24）工程机械指数分析旬度报告'
    
    # 合并修改和文档合并函数，执行文档修改
    merge_and_process_documents(input_folder, output_folder, perform_modification=True, use_filename_as_title=False)
