import os
import re
from docx import Document
from docxcompose.composer import Composer

def extract_numbers(s):
    return [int(match) for match in re.findall(r'\d+', s)]

def remove_second_paragraph(doc):
    if len(doc.paragraphs) >= 2:
        second_paragraph = doc.paragraphs[1]
        second_paragraph.clear()

        for run in second_paragraph.runs:
            new_run = doc.paragraphs[0].add_run(run.text, run.style)
            new_run.bold = run.bold

        doc._element.body.remove(second_paragraph._element)

def process_single_doc(doc, use_filename_as_title):
    first_paragraph = doc.paragraphs[0].text
    second_paragraph = doc.paragraphs[1].text

    if use_filename_as_title:
        first_paragraph = os.path.splitext(filename)[0]
    else:
        first_paragraph = f'{doc.paragraphs[0].text} {doc.paragraphs[1].text}'
        remove_second_paragraph(doc)

    doc.paragraphs[0].text = first_paragraph
    return first_paragraph, second_paragraph

def save_output_title(filename, first_paragraph, second_paragraph):
    with open('output_title.txt', 'a+') as f:
        f.write(f"{filename}: {first_paragraph}")
        f.write(f"{second_paragraph}\n")

def process_word_documents(input_folder, output_folder, use_filename_as_title=False):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
        print(f"Path {output_folder} created.")
    else:
        print(f"Path {output_folder} already exists.")

    if not os.path.exists(input_folder):
        print(f"Input folder '{input_folder}' does not exist. Exiting.")
        return

    file_list = os.listdir(input_folder)
    sorted_file_list = sorted(file_list, key=lambda x: extract_numbers(x))
    print("Sorted File List:")
    print("\n".join(sorted_file_list))

    with open("output.txt", "wt") as file:
        file.write(f"输出结果：\n")
        for i, filename in enumerate(sorted_file_list):
            if filename.endswith(".docx"):
                doc_path = os.path.join(input_folder, filename)
                doc = Document(doc_path)

                first_paragraph, second_paragraph = process_single_doc(doc, use_filename_as_title)

                print(filename, first_paragraph, second_paragraph, end="\n\n")
                output_filename = os.path.join(output_folder, filename)
                doc.save(output_filename)
                save_output_title(filename, first_paragraph, second_paragraph)

def merge_doc(source_file_path_list, target_folder, input_folder):
    page_break_doc = Document()
    page_break_doc.add_page_break()

    sorted_source_files = sorted(source_file_path_list, key=lambda x: extract_numbers(x))
    print("Sorted Source Files:")
    print("\n".join(sorted_source_files))

    input_folder_name = os.path.basename(os.path.normpath(input_folder))
    target_file_prefix = os.path.join(target_folder, f"{input_folder_name}合并")
    target_file_extension = ".docx"
    target_file_path = None

    target_doc = Document(sorted_source_files[0])
    target_composer = Composer(target_doc)
    for i in range(len(sorted_source_files)):
        if i == 0:
            continue
        target_composer.append(page_break_doc)
        f = sorted_source_files[i]
        target_composer.append(Document(f))

    target_file_path = f"{target_file_prefix}（1-{len(sorted_source_files)}）{target_file_extension}"
    target_composer.save(target_file_path)

def merge_and_process_documents(input_folder, output_folder, perform_modification=True, use_filename_as_title=False):
    if perform_modification:
        process_word_documents(input_folder, output_folder, use_filename_as_title)

    modified_files = os.listdir(output_folder)
    modified_files_paths = [os.path.join(output_folder, file) for file in modified_files]

    merge_doc(source_file_path_list=modified_files_paths, target_folder='/Users/harvin/code/docx_merge/data/merged', input_folder=input_folder)

if __name__ == '__main__':
    input_folder = '/Users/harvin/code/docx_merge/data/5.5 （25-24）工程机械指数分析旬度报告'
    output_folder = '/Users/harvin/code/docx_merge/data/output/5.5 （25-24）工程机械指数分析旬度报告'
    
    merge_and_process_documents(input_folder, output_folder, perform_modification=True, use_filename_as_title=False)
